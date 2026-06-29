from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body, BackgroundTasks
from sqlalchemy.orm import Session
from sqlalchemy import func, extract
from datetime import datetime, date, timedelta
from typing import List, Optional

from database import get_db
from security import get_teacher, get_current_user
import grading
from models import (
    User, TeacherProfile, ClassEntry, ClassStatus,
    RescheduleRequest, RescheduleStatus, DPP, Test, Doubt,
    DoubtStatus, Timetable, Notification, TestStatus,
    Exam, ExamQuestion, ExamAttempt, ExamResult
)
from schemas import (
    ClassEntryCreate, ClassEntryUpdate, ClassEntryOut,
    TimetableCreate, TimetableOut,
    RescheduleCreate, RescheduleOut,
    DPPCreate, DPPOut,
    TestCreate, TestPaperUpload, TestOut,
    DoubtResolve, DoubtOut,
    TeacherDashboard
)

router = APIRouter(prefix="/api/teacher", tags=["Teacher"])

def get_teacher_profile(user, db):
    profile = db.query(TeacherProfile).filter(TeacherProfile.user_id == user.id).first()
    if not profile:
        raise HTTPException(status_code=404, detail="Teacher profile nahi mila")
    return profile

def notify(db, user_id: int, title: str, message: str, notif_type: str):
    """Helper to create notification"""
    n = Notification(user_id=user_id, title=title, message=message, notif_type=notif_type)
    db.add(n)

# ===== DASHBOARD =====
@router.get("/dashboard", response_model=TeacherDashboard)
def teacher_dashboard(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    now = datetime.now()
    month_start = date(now.year, now.month, 1)
    week_start  = date.today() - timedelta(days=date.today().weekday())

    q = db.query(ClassEntry).filter(ClassEntry.teacher_id == tp.id)

    total_done      = q.filter(ClassEntry.status == ClassStatus.done).count()
    total_pending   = q.filter(ClassEntry.status == ClassStatus.pending).count()
    total_rescheduled = q.filter(ClassEntry.status == ClassStatus.rescheduled).count()
    monthly_done    = q.filter(ClassEntry.status == ClassStatus.done, ClassEntry.scheduled_date >= month_start).count()
    monthly_pending = q.filter(ClassEntry.status == ClassStatus.pending, ClassEntry.scheduled_date >= month_start).count()
    weekly_done     = q.filter(ClassEntry.status == ClassStatus.done, ClassEntry.scheduled_date >= week_start).count()

    # Reset monthly reschedule counter if new month
    if tp.reschedule_reset_month != now.month:
        tp.reschedule_count_this_month = 0
        tp.reschedule_reset_month = now.month
        db.commit()

    total_dpps  = db.query(DPP).filter(DPP.teacher_id == tp.id).count()
    total_tests = db.query(Test).filter(Test.teacher_id == tp.id).count()
    unresolved  = db.query(Doubt).filter(Doubt.teacher_id == tp.id, Doubt.status == DoubtStatus.pending).count()

    return TeacherDashboard(
        total_done=total_done, total_pending=total_pending,
        total_rescheduled=total_rescheduled, monthly_done=monthly_done,
        monthly_pending=monthly_pending, weekly_done=weekly_done,
        reschedule_this_month=tp.reschedule_count_this_month,
        total_dpps=total_dpps, total_tests=total_tests,
        unresolved_doubts=unresolved
    )

# ===== TIMETABLE =====
@router.post("/timetable", response_model=TimetableOut)
def add_timetable(req: TimetableCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    entry = Timetable(teacher_id=tp.id, **req.model_dump())
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry

@router.get("/timetable", response_model=List[TimetableOut])
def get_timetable(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(Timetable).filter(Timetable.teacher_id == tp.id, Timetable.is_active == True).all()

# ===== CLASSES =====
@router.post("/classes", response_model=ClassEntryOut)
def create_class(req: ClassEntryCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    entry = ClassEntry(teacher_id=tp.id, **req.model_dump())
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry

@router.get("/classes", response_model=List[ClassEntryOut])
def get_classes(
    status: Optional[str] = None,
    subject: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    q = db.query(ClassEntry).filter(ClassEntry.teacher_id == tp.id)
    if status:
        q = q.filter(ClassEntry.status == status)
    if subject:
        q = q.filter(ClassEntry.subject == subject)
    return q.order_by(ClassEntry.scheduled_date, ClassEntry.scheduled_time).all()

@router.get("/classes/today", response_model=List[ClassEntryOut])
def get_today_classes(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(ClassEntry).filter(
        ClassEntry.teacher_id == tp.id,
        ClassEntry.scheduled_date == date.today()
    ).order_by(ClassEntry.scheduled_time).all()

@router.patch("/classes/{class_id}/upload", response_model=ClassEntryOut)
def upload_class_pdf(
    class_id: int,
    drive_link: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    """Teacher uploads PDF → status auto-Done"""
    tp = get_teacher_profile(current_user, db)
    entry = db.query(ClassEntry).filter(
        ClassEntry.id == class_id,
        ClassEntry.teacher_id == tp.id
    ).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Class nahi mili")
    entry.drive_link = drive_link
    entry.status = ClassStatus.done
    db.commit()
    db.refresh(entry)
    return entry

# ===== RESCHEDULE =====
@router.post("/reschedule", response_model=RescheduleOut)
def request_reschedule(req: RescheduleCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    now = datetime.now()

    # Reset if new month
    if tp.reschedule_reset_month != now.month:
        tp.reschedule_count_this_month = 0
        tp.reschedule_reset_month = now.month

    # Check monthly limit
    if tp.reschedule_count_this_month >= 2:
        raise HTTPException(
            status_code=429,
            detail="LIMIT_REACHED: Aapne is mahine ki 2 reschedule limit poori kar li hai. Next month active hoga."
        )

    class_entry = db.query(ClassEntry).filter(
        ClassEntry.id == req.class_entry_id,
        ClassEntry.teacher_id == tp.id
    ).first()
    if not class_entry:
        raise HTTPException(status_code=404, detail="Class nahi mili")

    # Check existing pending request
    existing = db.query(RescheduleRequest).filter(
        RescheduleRequest.class_entry_id == req.class_entry_id,
        RescheduleRequest.status == RescheduleStatus.pending
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="Is class ke liye pehle se request pending hai")

    # Mark class as rescheduled (pending admin approval)
    class_entry.status = ClassStatus.rescheduled

    rs = RescheduleRequest(
        class_entry_id=req.class_entry_id,
        teacher_id=tp.id,
        original_date=class_entry.scheduled_date,
        original_time=class_entry.scheduled_time,
        new_date=req.new_date,
        new_time=req.new_time,
        reason=req.reason,
        status=RescheduleStatus.pending
    )
    db.add(rs)

    # Notify all admins
    admins = db.query(User).filter(User.role == "admin").all()
    for admin in admins:
        notify(db, admin.id,
               f"Reschedule Request — {current_user.name}",
               f"{class_entry.subject} ({class_entry.class_name}) ko {req.new_date} pe reschedule karna chahte hain. Reason: {req.reason}",
               "reschedule_request")

    db.commit()
    db.refresh(rs)
    return rs

@router.get("/reschedule", response_model=List[RescheduleOut])
def get_my_reschedules(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(RescheduleRequest).filter(RescheduleRequest.teacher_id == tp.id).all()

# ===== DPP =====
@router.post("/dpp", response_model=DPPOut)
def upload_dpp(req: DPPCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    dpp = DPP(teacher_id=tp.id, **req.model_dump())
    db.add(dpp)
    db.commit()
    db.refresh(dpp)
    return dpp

@router.get("/dpp", response_model=List[DPPOut])
def get_dpps(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(DPP).filter(DPP.teacher_id == tp.id).all()

# ===== TESTS =====
@router.post("/tests", response_model=TestOut)
def create_test(req: TestCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    test = Test(teacher_id=tp.id, **req.model_dump())
    db.add(test)
    db.commit()
    db.refresh(test)
    return test

@router.patch("/tests/{test_id}/upload-paper")
def upload_question_paper(
    test_id: int,
    drive_link: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    """Upload question paper — must be 15 min before test"""
    tp = get_teacher_profile(current_user, db)
    test = db.query(Test).filter(Test.id == test_id, Test.teacher_id == tp.id).first()
    if not test:
        raise HTTPException(status_code=404, detail="Test nahi mila")
    test.question_paper_link = drive_link
    test.status = TestStatus.active
    db.commit()
    return {"message": "Question paper upload ho gaya! Students ko access mil gayi.", "drive_link": drive_link}

@router.get("/tests", response_model=List[TestOut])
def get_tests(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(Test).filter(Test.teacher_id == tp.id).all()

# ===== DOUBTS =====
@router.get("/doubts")
def get_doubts(
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    q = db.query(Doubt).filter(Doubt.teacher_id == tp.id)
    if status:
        q = q.filter(Doubt.status == status)
    out = []
    for d in q.order_by(Doubt.created_at.desc()).all():
        sname = d.student.user.name if d.student and d.student.user else "Student"
        out.append({"id": d.id, "student_name": sname, "subject": d.subject, "topic": d.topic,
                    "question": d.question, "has_image": bool(d.image_b64),
                    "answer": d.answer, "status": d.status.value if hasattr(d.status, "value") else d.status,
                    "created_at": str(d.created_at)[:16]})
    return out

@router.get("/doubt/{did}/image")
def teacher_doubt_image(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    d = db.query(Doubt).filter(Doubt.id == did).first()
    if not d or not d.image_b64:
        raise HTTPException(status_code=404, detail="Image nahi")
    return Response(content=base64.b64decode(d.image_b64), media_type="image/jpeg")

@router.patch("/doubts/{doubt_id}/resolve")
def resolve_doubt(
    doubt_id: int,
    req: DoubtResolve,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    doubt = db.query(Doubt).filter(Doubt.id == doubt_id, Doubt.teacher_id == tp.id).first()
    if not doubt:
        raise HTTPException(status_code=404, detail="Doubt nahi mila")
    doubt.answer = req.answer
    doubt.answer_image_link = req.answer_image_link
    doubt.status = DoubtStatus.resolved
    doubt.resolved_at = datetime.now()

    # Notify student
    student_user = db.query(User).filter(User.id == doubt.student.user_id).first()
    if student_user:
        notify(db, student_user.id,
               "Aapka Doubt Resolve Ho Gaya! ✅",
               f"{doubt.subject} — {doubt.topic}: {current_user.name} ne jawab de diya.",
               "doubt_resolved")
    db.commit()
    return {"message": "Doubt resolve kar diya! Student ko notification chali gayi."}

# ===== NOTIFICATIONS =====
@router.get("/notifications")
def get_notifications(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    notifs = db.query(Notification).filter(
        Notification.user_id == current_user.id
    ).order_by(Notification.created_at.desc()).limit(20).all()
    return notifs

@router.patch("/notifications/{notif_id}/read")
def mark_read(notif_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    n = db.query(Notification).filter(Notification.id == notif_id, Notification.user_id == current_user.id).first()
    if n:
        n.is_read = True
        db.commit()
    return {"message": "Read mark ho gaya"}

# ===== TIMETABLE ENTRIES (chapter + parts + date + day) =====
def _serialize_tt(e):
    return {
        "id": e.id, "subject": e.subject, "class_name": e.class_name,
        "chapter": e.chapter, "part": e.part,
        "date": str(e.entry_date) if e.entry_date else None, "day": e.day,
        "time": getattr(e, "time_text", None), "type": getattr(e, "entry_type", None) or "chapter", "status": getattr(e, "status", None) or "approved",
        "completed": bool(getattr(e, "completed", False)),
        "topic_covered": getattr(e, "topic_covered", None),
        "homework": getattr(e, "homework", None),
        "dpp_given": bool(getattr(e, "dpp_given", False)),
        "remarks": getattr(e, "remarks", None),
        "start_time": getattr(e, "start_time", None),
        "end_time": getattr(e, "end_time", None),
        "completed_at": str(getattr(e, "completed_at", "")) if getattr(e, "completed_at", None) else None
    }

@router.post("/timetable-entry")
def add_tt_entry(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    edate = None
    d = (payload.get("entry_date") or "").strip()
    if d:
        try:
            edate = datetime.strptime(d, "%Y-%m-%d").date()
        except Exception:
            edate = None
    e = TimetableEntry(
        teacher_id=tp.id,
        subject=(payload.get("subject") or "").strip(),
        class_name=(payload.get("class_name") or "").strip(),
        chapter=(payload.get("chapter") or "").strip(),
        part=(payload.get("part") or "").strip() or None,
        entry_date=edate,
        day=(payload.get("day") or "").strip() or None,
        time_text=(payload.get("time") or "").strip() or None,
        entry_type=(payload.get("type") or "chapter").strip()
    )
    db.add(e); db.commit(); db.refresh(e)
    return _serialize_tt(e)

@router.get("/timetable-entries")
def list_tt_entries(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    es = db.query(TimetableEntry).filter(TimetableEntry.teacher_id == tp.id).order_by(
        TimetableEntry.subject, TimetableEntry.chapter, TimetableEntry.entry_date
    ).all()
    return [_serialize_tt(e) for e in es]

@router.delete("/timetable-entry/{entry_id}")
def delete_tt_entry(entry_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id, TimetableEntry.teacher_id == tp.id).first()
    if not e:
        raise HTTPException(status_code=404, detail="Entry nahi mili")
    db.delete(e); db.commit()
    return {"message": "Delete ho gaya"}

@router.delete("/timetable-entries/all")
def clear_tt_entries(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    db.query(TimetableEntry).filter(TimetableEntry.teacher_id == tp.id).delete()
    db.commit()
    return {"message": "Saari entries clear ho gayi"}

# ===== PDF TIMETABLE UPLOAD (auto-parse) =====
@router.post("/timetable-pdf")
async def upload_timetable_pdf(
    file: UploadFile = File(...),
    class_name: str = Form("Class 12"),
    subject: str = Form(""),
    replace: str = Form("false"),
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    import tt_parser
    raw = await file.read()
    try:
        rows = tt_parser.parse_pdf(raw, force_subject=(subject.strip() or None))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF parse error: {e}")
    if not rows:
        raise HTTPException(status_code=400, detail="PDF se koi valid row nahi mili. Text-based PDF honi chahiye.")

    subjects_found = sorted(set(r["subject"] for r in rows))
    if replace.lower() == "true":
        db.query(TimetableEntry).filter(
            TimetableEntry.teacher_id == tp.id,
            TimetableEntry.subject.in_(subjects_found)
        ).delete(synchronize_session=False)

    added = 0
    for r in rows:
        edate = None
        try:
            edate = datetime.strptime(r["date"], "%Y-%m-%d").date()
        except Exception:
            pass
        db.add(TimetableEntry(
            teacher_id=tp.id, subject=r["subject"], class_name=class_name,
            chapter=r["chapter"], part=r["part"], entry_date=edate,
            day=r["day"] or None, time_text=r["time"] or None, entry_type=r["type"]
        ))
        added += 1
    db.commit()
    return {"added": added, "subjects": subjects_found}

# ===== TEACHER: EDIT TIMETABLE ENTRY TOPIC/PART =====
@router.patch("/timetable-entry/{entry_id}")
def edit_tt_entry(entry_id: int, payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id).first()
    if not e or (e.subject not in (tp.subjects or []) and e.teacher_id != tp.id):
        raise HTTPException(status_code=404, detail="Entry nahi mili")
    if "part" in payload:
        e.part = (payload.get("part") or "").strip() or None
    if "time" in payload:
        e.time_text = (payload.get("time") or "").strip() or None
    db.commit()
    return _serialize_tt(e)

# ===== TEACHER: SEND NOTIFICATION TO STUDENTS =====
@router.post("/notify")
def teacher_notify(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import User
    title = (payload.get("title") or "").strip()
    message = (payload.get("message") or "").strip()
    if not title or not message:
        raise HTTPException(status_code=400, detail="Title aur message zaroori hain")
    students = db.query(User).filter(User.is_active == True, User.role == "student").all()
    sender = "👨‍🏫 " + current_user.name
    for s in students:
        notify(db, s.id, sender + ": " + title, message, "teacher_message")
    db.commit()
    return {"message": f"{len(students)} students ko bhej di!", "count": len(students)}

# ===== STUDY MATERIAL (PDF upload to DB) =====
@router.post("/material")
async def upload_material(
    file: UploadFile = File(...),
    subject: str = Form(...),
    class_name: str = Form("Class 12"),
    chapter: str = Form(""),
    material_type: str = Form("notes"),   # notes | dpp | test | other
    title: str = Form(""),
    category: str = Form(""),
    duration_min: int = Form(0),
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    import base64
    from models import Material
    tp = get_teacher_profile(current_user, db)
    raw = await file.read()
    if len(raw) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File 20MB se badi hai. Chhoti PDF use karein.")
    b64 = base64.b64encode(raw).decode("ascii")
    m = Material(
        teacher_id=tp.id, teacher_name=current_user.name, subject=subject.strip(),
        class_name=class_name.strip(), chapter=chapter.strip(),
        material_type=material_type.strip(), title=(title.strip() or file.filename),
        category=(category.strip() or None),
        filename=file.filename, content_b64=b64,
        duration_min=(duration_min or None)
    )
    db.add(m); db.commit(); db.refresh(m)
    # Notify students who have this subject
    try:
        from models import StudentProfile
        label = {"notes": "Class Notes", "dpp": "DPP", "test": "Test"}.get(m.material_type, (m.category or "Material"))
        sps = db.query(StudentProfile).all()
        for sp in sps:
            if sp.subjects and subject.strip() in sp.subjects and sp.user:
                notify(db, sp.user.id, f"📚 New {label}: {subject.strip()}",
                       f"{current_user.name} ne {subject.strip()} ({chapter.strip() or 'General'}) ke liye {label} upload ki hai. Materials section mein dekho!",
                       "new_material")
        db.commit()
    except Exception:
        db.rollback()
    return {"id": m.id, "message": "Upload ho gaya!"}

@router.get("/materials")
def teacher_materials(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    tp = get_teacher_profile(current_user, db)
    ms = db.query(Material).filter(Material.teacher_id == tp.id,
                                   Material.material_type != "answer").order_by(Material.created_at.desc()).all()
    return [{"id": m.id, "subject": m.subject, "chapter": m.chapter, "type": m.material_type,
             "title": m.title, "filename": m.filename, "duration_min": m.duration_min,
             "date": str(m.created_at)[:10]} for m in ms]

@router.get("/chapter-status")
def chapter_status(subject: str, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """For a subject, list chapters from timetable + whether notes/dpp uploaded."""
    from models import TimetableEntry, Material
    chapters = [r[0] for r in db.query(TimetableEntry.chapter).filter(
        TimetableEntry.subject == subject,
        TimetableEntry.entry_type == "chapter").distinct().all() if r[0]]
    mats = db.query(Material).filter(Material.subject == subject).all()
    out = []
    for ch in chapters:
        notes = any(m.chapter == ch and m.material_type == "notes" for m in mats)
        dpp = any(m.chapter == ch and m.material_type == "dpp" for m in mats)
        out.append({"chapter": ch, "notes": notes, "dpp": dpp})
    return out

@router.get("/material/{mid}/download")
def teacher_download(mid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    from models import Material
    m = db.query(Material).filter(Material.id == mid).first()
    if not m: raise HTTPException(status_code=404, detail="Nahi mila")
    data = base64.b64decode(m.content_b64)
    return Response(content=data, media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{m.filename or "file.pdf"}"'})

@router.delete("/material/{mid}")
def delete_material(mid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    tp = get_teacher_profile(current_user, db)
    m = db.query(Material).filter(Material.id == mid, Material.teacher_id == tp.id).first()
    if not m: raise HTTPException(status_code=404, detail="Nahi mila")
    db.delete(m); db.commit()
    return {"message": "Delete ho gaya"}

# ===== TEACHER PROFILE & SUBJECT SELECTION (class-wise) =====
@router.get("/profile")
def teacher_profile(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    sc = tp.subject_classes or []
    return {
        "name": current_user.name,
        "user_id": current_user.user_id,
        "gender": tp.gender,
        "subjects": tp.subjects or [],
        "subject_classes": sc,
        "needs_subjects": len(sc) == 0
    }

@router.get("/available-subjects")
def teacher_available_subjects(class_level: str, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import AvailableSubject
    subs = db.query(AvailableSubject).filter(
        AvailableSubject.class_level == class_level, AvailableSubject.is_active == True).all()
    return [{"name": s.name, "code": s.code} for s in subs]

@router.post("/set-subjects")
def teacher_set_subjects(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    selections = payload.get("selections", [])   # [{"subject":..,"class":"10"/"12"}]
    if not selections:
        raise HTTPException(status_code=400, detail="Kam se kam 1 subject select karein")
    tp.subject_classes = selections
    tp.subjects = sorted({s.get("subject") for s in selections if s.get("subject")})
    db.commit()
    return {"message": "Subjects save ho gaye!", "subjects": tp.subjects}

# ===== TEACHER: VIEW TIMETABLE (by their subjects, admin-uploaded) =====
@router.get("/my-timetable")
def my_timetable(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    from sqlalchemy import or_
    es = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status==None, TimetableEntry.status!='pending')).order_by(
        TimetableEntry.subject, TimetableEntry.entry_date).all()
    return [_serialize_tt(e) for e in es]

# ===== TEACHER: TODAY'S CLASSES with material status =====
@router.get("/today-classes")
def today_classes(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    today = date.today()
    from sqlalchemy import or_
    es = db.query(TimetableEntry).filter(
        TimetableEntry.subject.in_(subs), TimetableEntry.entry_date == today,
        or_(TimetableEntry.status==None, TimetableEntry.status!='pending')).all()
    mats = db.query(Material).filter(Material.subject.in_(subs)).all()
    out = []
    for e in es:
        notes = any(m.chapter == e.chapter and m.subject == e.subject and m.material_type == "notes" for m in mats)
        dpp = any(m.chapter == e.chapter and m.subject == e.subject and m.material_type == "dpp" for m in mats)
        d = _serialize_tt(e); d["notes"] = notes; d["dpp"] = dpp
        out.append(d)
    out.sort(key=lambda x: x.get("time") or "")
    return out

# ===== TEACHER: REQUEST EXTRA CLASS (needs admin approval) =====
@router.post("/request-class")
def request_class(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, User, UserRole, Notification
    tp = get_teacher_profile(current_user, db)
    subject = (payload.get("subject") or "").strip()
    if subject not in (tp.subjects or []):
        raise HTTPException(status_code=400, detail="Yeh aapka subject nahi hai")
    edate = None
    if payload.get("date"):
        try:
            from datetime import datetime as _dt
            edate = _dt.strptime(payload["date"], "%Y-%m-%d").date()
        except Exception:
            pass
    day = edate.strftime("%a") if edate else None
    e = TimetableEntry(
        teacher_id=tp.id, subject=subject, class_name=payload.get("class_name", "Class 12"),
        chapter=(payload.get("topic") or "Extra Class").strip(), part=None,
        entry_date=edate, day=day, time_text=(payload.get("time") or "").strip() or None,
        entry_type="chapter", status="pending"
    )
    db.add(e); db.flush()
    # notify admins
    for adm in db.query(User).filter(User.role == UserRole.admin).all():
        db.add(Notification(user_id=adm.id, title="New Extra Class Request",
                            message=f"{current_user.name} ne {subject} ki extra class request ki hai ({payload.get('date','')} {payload.get('time','')}). Approve karein.",
                            notif_type="class_request"))
    db.commit(); db.refresh(e)
    return {"id": e.id, "message": "Request admin ko bhej di! Approve hote hi timetable mein aa jayegi."}

# ===== TEACHER: SUBJECT-WISE STUDENT COUNTS =====
@router.get("/student-counts")
def teacher_student_counts(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import StudentProfile
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    sc = tp.subject_classes or []
    students = db.query(StudentProfile).all()
    out = []
    for s in subs:
        cnt = sum(1 for sp in students if sp.subjects and s in sp.subjects)
        cls = next((x.get("class") for x in sc if x.get("subject") == s), None)
        out.append({"subject": s, "class": cls, "count": cnt})
    out.sort(key=lambda x: -x["count"])
    return {"total": sum(o["count"] for o in out), "subjects": out}

# ===== TEACHER: VIEW SUBMISSIONS + GIVE MARKS =====
@router.get("/submissions")
def teacher_submissions(parent_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    subs = db.query(Material).filter(Material.material_type == "answer",
                                     Material.parent_id == parent_id).order_by(Material.created_at.desc()).all()
    return [{"id": m.id, "student_name": m.student_name, "marks": m.marks,
             "date": str(m.created_at)[:16]} for m in subs]

@router.post("/submission/{sid}/marks")
def set_marks(sid: int, payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, StudentProfile, Notification
    m = db.query(Material).filter(Material.id == sid, Material.material_type == "answer").first()
    if not m:
        raise HTTPException(status_code=404, detail="Submission nahi mili")
    m.marks = str(payload.get("marks", "")).strip()
    # notify student
    if m.student_id:
        sp = db.query(StudentProfile).filter(StudentProfile.id == m.student_id).first()
        if sp and sp.user:
            db.add(Notification(user_id=sp.user.id, title="DPP Checked!",
                                message=f"{current_user.name} ne aapki {m.subject} DPP check ki. Marks: {m.marks}",
                                notif_type="marks"))
    db.commit()
    return {"message": "Marks save ho gaye!"}

# ===== TEACHER: OWN PHOTO + MY STUDENTS LIST =====
@router.get("/my-photo")
def teacher_my_photo(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    if not tp.photo_b64:
        raise HTTPException(status_code=404, detail="Photo nahi")
    return Response(content=base64.b64decode(tp.photo_b64), media_type="image/jpeg")

@router.get("/student/{sid}/photo")
def teacher_student_photo(sid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    from models import StudentProfile
    sp = db.query(StudentProfile).filter(StudentProfile.id == sid).first()
    if not sp or not sp.photo_b64:
        raise HTTPException(status_code=404, detail="Photo nahi")
    return Response(content=base64.b64decode(sp.photo_b64), media_type="image/jpeg")

@router.get("/my-students-list")
def teacher_my_students_list(q: str = "", subject: str = "", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import StudentProfile
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    ql = q.strip().lower()
    rows = db.query(StudentProfile).all()
    out = []
    for sp in rows:
        ssubs = sp.subjects or []
        if not any(s in subs for s in ssubs):
            continue
        if subject and subject not in ssubs:
            continue
        nm = sp.user.name if sp.user else ""
        if ql and ql not in nm.lower() and ql not in (sp.phone or ""):
            continue
        out.append({"id": sp.id, "name": nm, "phone": sp.phone, "class": sp.class_level,
                    "subjects": [s for s in ssubs if s in subs], "has_photo": bool(sp.photo_b64)})
    out.sort(key=lambda x: x["name"].lower())
    return {"total": len(out), "students": out}

# ===== TEACHER -> ADMIN MESSAGE =====
@router.post("/message-admin")
def teacher_message_admin(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    title = (payload.get("title") or "").strip()
    message = (payload.get("message") or "").strip()
    if not title or not message:
        raise HTTPException(status_code=400, detail="Title and message are required")
    admins = db.query(User).filter(User.role == "admin").all()
    sender = current_user.name
    for a in admins:
        notify(db, a.id, f"\u2709\ufe0f {sender}: {title}", message, "teacher_to_admin")
    db.commit()
    return {"message": "Message sent to the admin"}

# ===== TEACHER ACCOUNTABILITY: classes with status, mark-complete, compliance =====
def _class_status(e):
    """Upcoming | Live | Completed | Missed based on date/time + completed flag."""
    if getattr(e, "completed", False):
        return "Completed"
    today = date.today()
    if e.entry_date is None:
        return "Upcoming"
    if e.entry_date < today:
        return "Missed"
    if e.entry_date > today:
        return "Upcoming"
    return "Pending"  # today, not yet completed

@router.get("/my-classes")
def teacher_my_classes(scope: str = "all", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    q = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"))
    if scope == "today":
        q = q.filter(TimetableEntry.entry_date == date.today())
    es = q.order_by(TimetableEntry.entry_date, TimetableEntry.time_text).all()
    out = []
    for e in es:
        d = _serialize_tt(e); d["live_status"] = _class_status(e)
        out.append(d)
    return out

@router.post("/class/{entry_id}/complete")
def teacher_complete_class(entry_id: int, payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    tp = get_teacher_profile(current_user, db)
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id).first()
    if not e or e.subject not in (tp.subjects or []):
        raise HTTPException(status_code=404, detail="Class not found")
    e.completed = True
    e.completed_at = datetime.now()
    e.topic_covered = (payload.get("topic_covered") or e.chapter or "").strip() or None
    e.start_time = (payload.get("start_time") or "").strip() or None
    e.end_time = (payload.get("end_time") or "").strip() or None
    e.homework = (payload.get("homework") or "").strip() or None
    e.dpp_given = bool(payload.get("dpp_given"))
    e.remarks = (payload.get("remarks") or "").strip() or None
    db.commit()
    return {"message": "Class marked as completed."}

@router.get("/compliance")
def teacher_compliance(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    subj_count = max(1, len(subs))
    today = date.today()
    month_start = date(today.year, today.month, 1)
    classes = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"),
        TimetableEntry.entry_type == "chapter").all() if subs else []
    due = [c for c in classes if c.entry_date and c.entry_date <= today]
    completed = [c for c in due if getattr(c, "completed", False)]
    mats = db.query(Material).filter(Material.subject.in_(subs)).all() if subs else []
    dpp_count = sum(1 for m in mats if m.material_type == "dpp")
    notes_count = sum(1 for m in mats if m.material_type == "notes")
    test_count = sum(1 for m in mats if m.material_type == "test")
    # component scores (0..1)
    cc = (len(completed) / len(due)) if due else 1.0
    dpp_s = min(1.0, dpp_count / subj_count)
    mat_s = min(1.0, notes_count / subj_count)
    test_s = min(1.0, test_count / subj_count)
    score = round(cc * 40 + dpp_s * 25 + mat_s * 20 + test_s * 15)
    band = "green" if score >= 81 else ("yellow" if score >= 61 else "red")
    return {
        "score": score, "band": band,
        "breakdown": {
            "class_completion": {"weight": 40, "pct": round(cc * 100), "got": round(cc * 40)},
            "dpp_upload": {"weight": 25, "pct": round(dpp_s * 100), "got": round(dpp_s * 25)},
            "study_material": {"weight": 20, "pct": round(mat_s * 100), "got": round(mat_s * 20)},
            "test_creation": {"weight": 15, "pct": round(test_s * 100), "got": round(test_s * 15)},
        },
        "stats": {
            "classes_due": len(due), "classes_completed": len(completed),
            "dpp_count": dpp_count, "notes_count": notes_count, "test_count": test_count,
            "classes_today": sum(1 for c in classes if c.entry_date == today),
            "completed_today": sum(1 for c in completed if c.entry_date == today),
            "pending_today": sum(1 for c in classes if c.entry_date == today and not getattr(c, "completed", False)),
            "missed": sum(1 for c in due if not getattr(c, "completed", False)),
            "subject_count": len(subs),
        }
    }

# ===== TEACHER: DOUBT STATS (pending, resolved, avg response time) =====
@router.get("/doubt-stats")
def teacher_doubt_stats(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Doubt, DoubtStatus
    tp = get_teacher_profile(current_user, db)
    ds = db.query(Doubt).filter(Doubt.teacher_id == tp.id).all()
    pending = sum(1 for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "pending")
    resolved_list = [d for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved" and d.resolved_at and d.created_at]
    resolved = sum(1 for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved")
    avg_min = None
    if resolved_list:
        total = sum((d.resolved_at - d.created_at).total_seconds() for d in resolved_list)
        avg_min = round(total / len(resolved_list) / 60)
    return {"pending": pending, "resolved": resolved, "total": len(ds), "avg_response_minutes": avg_min}

# ===== TEACHER: PERFORMANCE (aggregates + recent activity + monthly) =====
@router.get("/performance")
def teacher_performance(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    classes = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"),
        TimetableEntry.entry_type == "chapter").all() if subs else []
    completed = [c for c in classes if getattr(c, "completed", False)]
    mats = db.query(Material).filter(Material.subject.in_(subs)).all() if subs else []
    dpp_count = sum(1 for m in mats if m.material_type == "dpp")
    notes_count = sum(1 for m in mats if m.material_type in ("notes", "other"))
    test_count = sum(1 for m in mats if m.material_type == "test")
    # monthly completed (last 6 months)
    from collections import OrderedDict
    today = date.today()
    months = OrderedDict()
    for i in range(5, -1, -1):
        y = today.year; mo = today.month - i
        while mo <= 0:
            mo += 12; y -= 1
        months[f"{y}-{mo:02d}"] = 0
    for c in completed:
        if c.completed_at:
            key = f"{c.completed_at.year}-{c.completed_at.month:02d}"
            if key in months:
                months[key] += 1
    monthly = [{"month": k, "count": v} for k, v in months.items()]
    # recent activity (completions + uploads)
    acts = []
    for c in completed:
        if c.completed_at:
            acts.append({"type": "class", "text": f"Completed {c.subject} — {c.topic_covered or c.chapter or ''}", "at": c.completed_at})
    for m in mats:
        if m.created_at:
            acts.append({"type": m.material_type, "text": f"Uploaded {m.material_type.upper()}: {m.title or m.chapter or m.subject}", "at": m.created_at})
    acts.sort(key=lambda x: x["at"], reverse=True)
    recent = [{"type": a["type"], "text": a["text"], "at": str(a["at"])[:16]} for a in acts[:12]]
    return {
        "classes_assigned": len(classes), "classes_completed": len(completed),
        "dpp_uploaded": dpp_count, "materials_uploaded": notes_count, "tests_created": test_count,
        "monthly": monthly, "recent": recent
    }

# ===== TEACHER: MATERIAL ANALYTICS (views/downloads per material) =====
@router.get("/material-analytics")
def teacher_material_analytics(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, MaterialView
    from sqlalchemy import func as _f
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    mats = db.query(Material).filter(Material.subject.in_(subs),
        Material.material_type.in_(["notes", "dpp", "test", "other"])).order_by(Material.created_at.desc()).all() if subs else []
    out = []
    for m in mats:
        viewed = db.query(_f.count(_f.distinct(MaterialView.student_id))).filter(MaterialView.material_id == m.id).scalar() or 0
        downloads = db.query(_f.count(MaterialView.id)).filter(MaterialView.material_id == m.id, MaterialView.action == "download").scalar() or 0
        out.append({
            "id": m.id, "type": m.material_type, "category": m.category,
            "title": m.title or m.chapter or m.subject, "subject": m.subject,
            "upload_date": str(m.created_at)[:10] if m.created_at else None,
            "students_viewed": viewed, "downloads": downloads,
            "approval_status": getattr(m, "approval_status", "approved") or "approved",
        })
    return out

# ===== TEACHER: STUDENT ENGAGEMENT =====
@router.get("/student-engagement")
def teacher_student_engagement(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, MaterialView, StudentProfile
    from sqlalchemy import func as _f, or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    # students who have any of the teacher's subjects
    students = []
    for sp in db.query(StudentProfile).all():
        if set(sp.subjects or []) & set(subs):
            students.append(sp)
    # teacher material ids
    mat_ids = [m.id for m in db.query(Material).filter(Material.subject.in_(subs)).all()]
    out = []
    for sp in students:
        answers = db.query(Material).filter(Material.material_type == "answer", Material.student_id == sp.id).all()
        pids = [a.parent_id for a in answers if a.parent_id]
        ptypes = {}
        if pids:
            for pm in db.query(Material).filter(Material.id.in_(pids)).all():
                ptypes[pm.id] = pm.material_type
        dpp_done = sum(1 for a in answers if ptypes.get(a.parent_id) == "dpp")
        test_done = sum(1 for a in answers if ptypes.get(a.parent_id) == "test")
        downloads = db.query(_f.count(MaterialView.id)).filter(
            MaterialView.student_id == sp.id, MaterialView.action == "download",
            MaterialView.material_id.in_(mat_ids) if mat_ids else False).scalar() or 0
        last_act = db.query(MaterialView).filter(MaterialView.student_id == sp.id).order_by(MaterialView.created_at.desc()).first()
        out.append({
            "name": (sp.user.name if sp.user else "Student"),
            "phone": sp.phone, "subjects": sp.subjects or [],
            "dpp_completed": dpp_done, "tests_completed": test_done,
            "material_downloads": downloads,
            "last_active": str(last_act.created_at)[:16] if last_act else None,
        })
    out.sort(key=lambda x: (x["material_downloads"] + x["dpp_completed"] + x["tests_completed"]), reverse=True)
    return out


# ===================== EXAM / TEST ENGINE (teacher) =====================
@router.post("/exam")
def create_exam(payload: dict = Body(...), background_tasks: BackgroundTasks = None, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    qs = payload.get("questions") or []
    if not payload.get("title") or not qs:
        raise HTTPException(400, "Title and at least one question are required")
    ttype = payload.get("test_type", "subjective")
    total = sum(int(q.get("max_marks", 1) or 1) for q in qs)
    ex = Exam(teacher_id=tp.id, teacher_name=current_user.name,
              subject=payload.get("subject", ""), title=payload["title"],
              chapter=payload.get("chapter"), test_type=ttype,
              medium=payload.get("medium", "English"),
              total_marks=total, duration_min=int(payload.get("duration_min", 60) or 60))
    db.add(ex); db.flush()
    for i, q in enumerate(qs, start=1):
        co = q.get("correct_option")
        opts_hi = q.get("options_hi") if ttype == "mcq" else None
        db.add(ExamQuestion(exam_id=ex.id, q_no=i,
               question_text=q.get("question_text", ""),
               max_marks=int(q.get("max_marks", 1) or 1),
               model_answer=q.get("model_answer"),
               options=q.get("options") if ttype == "mcq" else None,
               correct_option=(str(co) if co not in (None, "") else None),
               image_b64=q.get("image_b64"),
               question_text_hi=(q.get("question_text_hi") or None),
               model_answer_hi=(q.get("model_answer_hi") or None),
               options_hi=(opts_hi if opts_hi else None),
               model_answer_image=q.get("model_answer_image")))
    db.commit()
    # Bilingual test: auto-translate any blank Hindi fields in the background.
    if (ex.medium or "").lower().startswith("bi") and background_tasks is not None:
        background_tasks.add_task(_bg_translate_exam, ex.id)
    return {"id": ex.id, "total_marks": total, "questions": len(qs),
            "test_type": ttype, "medium": ex.medium}


def _bg_translate_exam(exam_id):
    """Fill in any missing Hindi fields for a bilingual test using Gemini.
    Runs after the response so test creation stays fast. Only fills blanks."""
    from database import SessionLocal
    db = SessionLocal()
    try:
        ex = db.query(Exam).filter(Exam.id == exam_id).first()
        if not ex:
            return
        qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
        for q in qs:
            need_q = not (q.question_text_hi or "").strip()
            need_a = (ex.test_type != "mcq") and not (q.model_answer_hi or "").strip()
            need_o = (ex.test_type == "mcq") and not q.options_hi
            if not (need_q or need_a or need_o):
                continue
            tr = grading.translate_question_to_hindi(
                q.question_text or "", q.model_answer or "",
                (q.options or []) if ex.test_type == "mcq" else None, ex.subject or "")
            if not tr:
                continue
            if need_q and tr.get("question"):
                q.question_text_hi = tr["question"]
            if need_a and tr.get("answer"):
                q.model_answer_hi = tr["answer"]
            if need_o and tr.get("options") and len(tr["options"]) == len(q.options or []):
                q.options_hi = tr["options"]
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()


@router.get("/exam/{exam_id}/pdf")
def teacher_exam_pdf(exam_id: int, medium: str = "english",
                     db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Download the full question+answer paper as a PDF in English or Hindi medium."""
    import exam_pdf
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Test not found")
    qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
    med = "hindi" if str(medium).lower().startswith("hi") else "english"
    try:
        data = exam_pdf.build_exam_pdf(ex, qs, med)
    except Exception as e:
        raise HTTPException(500, "Could not generate the PDF. The server needs fpdf2, "
                                 "uharfbuzz and the Devanagari font. (%s)" % e)
    safe = (ex.title or "test").replace('"', "").replace("/", "-").strip()[:60] or "test"
    return Response(content=data, media_type="application/pdf",
                    headers={"Content-Disposition": 'attachment; filename="%s-%s.pdf"' % (safe, med)})

@router.get("/exams")
def list_exams(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    rows = db.query(Exam).filter(Exam.teacher_id == tp.id, Exam.is_active == True).order_by(Exam.created_at.desc()).all()
    out = []
    for e in rows:
        nq = db.query(ExamQuestion).filter(ExamQuestion.exam_id == e.id).count()
        na = db.query(ExamAttempt).filter(ExamAttempt.exam_id == e.id).count()
        ng = db.query(ExamAttempt).filter(ExamAttempt.exam_id == e.id, ExamAttempt.status == "graded").count()
        out.append({"id": e.id, "title": e.title, "subject": e.subject, "chapter": e.chapter,
                    "test_type": e.test_type, "total_marks": e.total_marks, "duration_min": e.duration_min,
                    "medium": e.medium, "questions": nq, "attempts": na, "graded": ng,
                    "created_at": e.created_at.isoformat() if e.created_at else None})
    return out

@router.get("/exam/{exam_id}/attempts")
def exam_attempts(exam_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Exam not found")
    atts = db.query(ExamAttempt).filter(ExamAttempt.exam_id == exam_id).order_by(ExamAttempt.submitted_at.desc()).all()
    out = [{"attempt_id": a.id, "student_id": a.student_id, "student_name": a.student_name,
            "status": a.status, "total_awarded": a.total_awarded, "verdict": a.verdict,
            "submitted_at": a.submitted_at.isoformat() if a.submitted_at else None} for a in atts]
    qrows = db.query(ExamQuestion).filter(ExamQuestion.exam_id == exam_id).order_by(ExamQuestion.q_no).all()
    questions = [{"q_no": q.q_no, "question_text": q.question_text, "max_marks": q.max_marks} for q in qrows]
    return {"exam": {"id": ex.id, "title": ex.title, "total_marks": ex.total_marks, "test_type": ex.test_type},
            "questions": questions, "attempts": out}


def _exam_verdict_t(aw, tot):
    if not tot:
        return "Good"
    p = aw / tot * 100
    return "Excellent" if p >= 80 else ("Good" if p >= 50 else "Needs Improvement")

def _notify_exam_result_t(db, att, ex):
    """Notify the student that their test result is ready."""
    try:
        from models import StudentProfile
        sp = db.query(StudentProfile).filter(StudentProfile.id == att.student_id).first()
        if sp and sp.user_id:
            try:
                sc = "%g" % float(att.total_awarded)
            except Exception:
                sc = str(att.total_awarded)
            db.add(Notification(
                user_id=sp.user_id,
                title="Result ready: %s" % (ex.title or "Test"),
                message="Your test has been checked. You scored %s/%s. Tap to view your result and download your answer sheet." % (sc, ex.total_marks),
                notif_type="exam_result"))
    except Exception:
        pass


@router.post("/attempt/{attempt_id}/grade")
def grade_attempt_now(attempt_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    if att.status == "graded":
        return {"status": "graded", "message": "Already graded"}
    qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
    teacher = ex.teacher_name or "your teacher"
    if ex.test_type == "mcq":
        results, total = grading.grade_mcq(qs, att.mcq_answers or {})
        feedback, verdict = "", _exam_verdict_t(total, ex.total_marks)
    else:
        results, total, feedback, verdict = grading.grade_subjective(qs, att.answer_image_b64 or "", "image/jpeg")
        if results is None:
            raise HTTPException(400, "AI grading failed: " + (feedback or "unknown error") + " -- you can use Grade Manually instead.")
        verdict = verdict or _exam_verdict_t(total, ex.total_marks)
    db.query(ExamResult).filter(ExamResult.attempt_id == att.id).delete()
    for r in results:
        db.add(ExamResult(attempt_id=att.id, q_no=r["q_no"], marks_awarded=r["marks"],
               max_marks=r["max"], remark=r.get("remark", "")))
    att.total_awarded = total
    att.status = "graded"
    att.graded_at = datetime.utcnow()
    att.verdict = verdict
    att.overall_feedback = feedback or ("Graded by teacher. \u2014 %s" % teacher)
    _notify_exam_result_t(db, att, ex)
    db.commit()
    return {"status": "graded", "total_awarded": total, "verdict": verdict}


# ===================== AI AUTO-MAGIC ENDPOINTS (Phase 2) =====================
@router.post("/ocr-question")
def ocr_question(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    img = payload.get("image_b64") or ""
    if not img:
        raise HTTPException(400, "No image provided")
    res = grading.ocr_extract_question(img, payload.get("test_type", "subjective"),
                                       payload.get("mime_type", "image/jpeg"))
    if res is None:
        raise HTTPException(503, "AI could not read the image. Check GEMINI_API_KEY or try a clearer screenshot.")
    return res

@router.post("/format-text")
def format_text(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    out = grading.format_text_latex(payload.get("text", ""))
    if out is None:
        raise HTTPException(503, "AI formatting is unavailable. Check GEMINI_API_KEY.")
    return {"text": out}

@router.post("/parse-exam-docx")
async def parse_exam_docx(file: UploadFile = File(...), test_type: str = Form("subjective"),
                          db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import io
    try:
        from docx import Document
    except Exception:
        raise HTTPException(503, "Word parsing is not enabled on the server (add python-docx to requirements.txt).")
    data = await file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        raise HTTPException(400, "Could not open the Word file. Please upload a valid .docx file.")
    full = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
    # also pull text from tables
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                full += "\n" + " | ".join(cells)
    if not full.strip():
        raise HTTPException(400, "The Word file appears to be empty.")
    qs = grading.structure_docx_questions(full, test_type)
    if qs is None:
        raise HTTPException(503, "AI could not structure the document. Check GEMINI_API_KEY.")
    return {"questions": qs, "count": len(qs)}


@router.get("/attempt/{attempt_id}/answer")
def attempt_answer_image(attempt_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    if not att.answer_image_b64:
        raise HTTPException(404, "No answer sheet uploaded")
    raw = att.answer_image_b64.split(",")[-1]
    return Response(content=base64.b64decode(raw), media_type="image/jpeg")

@router.post("/attempt/{attempt_id}/grade-manual")
def grade_attempt_manual(attempt_id: int, payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    qmap = {q.q_no: q for q in db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).all()}
    results = payload.get("results") or []
    db.query(ExamResult).filter(ExamResult.attempt_id == att.id).delete()
    total = 0.0
    for r in results:
        try:
            qn = int(r.get("q_no"))
        except Exception:
            continue
        mx = qmap[qn].max_marks if qn in qmap else int(r.get("max", 1) or 1)
        try:
            mk = float(r.get("marks", 0) or 0)
        except Exception:
            mk = 0.0
        mk = max(0.0, min(mk, float(mx)))
        total += mk
        db.add(ExamResult(attempt_id=att.id, q_no=qn, marks_awarded=mk, max_marks=mx, remark=r.get("remark", "")))
    att.total_awarded = total
    att.status = "graded"
    att.graded_at = datetime.utcnow()
    att.verdict = payload.get("verdict") or _exam_verdict_t(total, ex.total_marks)
    fb = payload.get("feedback") or ""
    att.overall_feedback = fb if fb else ("Checked by %s." % (ex.teacher_name or "your teacher"))
    _notify_exam_result_t(db, att, ex)
    db.commit()
    return {"status": "graded", "total_awarded": total, "verdict": att.verdict}


@router.get("/ai-status")
def ai_status(current_user=Depends(get_teacher)):
    return grading.ai_status()
